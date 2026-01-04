from sync_fetch_data import initialize_and_fetch_db, get_empty_doc_record, upsert_question_record
from transform_question_to_vector import simplify_question_record
from data_utils import load_image
from transformers import BlipProcessor, BlipForConditionalGeneration
from PIL import Image
import torch
import pytesseract
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from pylatexenc import macrospec, latexwalker
from pylatexenc.latexwalker import LatexMacroNode, LatexGroupNode, LatexCharsNode
import os
import sys
import psutil

def convert_latex_to_plain_english(doc: str) -> str:
    """
    Converts LaTeX mathematical expressions to plain English.
    Finds LaTeX delimited by $ or $$ and replaces with English equivalents.
    Skips chemical notation.
    """
    lwc = latexwalker.get_default_latex_context_db()
    lwc.add_context_category('powers', specials=[
        macrospec.SpecialsSpec('^', args_parser=macrospec.MacroStandardArgsParser('{')),
        macrospec.SpecialsSpec('_', args_parser=macrospec.MacroStandardArgsParser('{')),
    ])
    
    def process_node(node):
        """Recursively process a single LaTeX node to English."""
        if isinstance(node, LatexCharsNode):
            text = node.chars
            text = text.replace('+', ' plus ')
            text = text.replace('-', ' minus ')
            text = text.replace('=', ' equals ')
            text = text.replace('/', ' divided by ')
            text = text.replace('*', ' times ')
            text = text.replace('%', ' percent ')
            text = text.replace('<', ' less than ')
            text = text.replace('>', ' greater than ')
            text = text.replace('(', ' open parenthesis ')
            text = text.replace(')', ' close parenthesis ')
            return text
        
        if isinstance(node, LatexGroupNode):
            return ''.join(process_node(n) for n in node.nodelist)
        
        if isinstance(node, latexwalker.LatexSpecialsNode):
            if node.specials_chars == '^':
                if node.nodeargd and hasattr(node.nodeargd, 'argnlist') and len(node.nodeargd.argnlist) > 0:
                    power = process_node(node.nodeargd.argnlist[0])
                    return f' to the power of {power}'
                return ' to the power of '
            elif node.specials_chars == '_':
                if node.nodeargd and hasattr(node.nodeargd, 'argnlist') and len(node.nodeargd.argnlist) > 0:
                    sub = process_node(node.nodeargd.argnlist[0])
                    return f' subscript {sub}'
                return ' subscript '
        
        if isinstance(node, LatexMacroNode):
            macro = node.macroname
            
            if macro == 'int':
                return ' integral '
            elif macro == 'sum':
                return ' summation '
            elif macro == 'prod':
                return ' product '
            elif macro == 'lim':
                return ' limit '
            elif macro == 'frac':
                if node.nodeargd and hasattr(node.nodeargd, 'argnlist') and len(node.nodeargd.argnlist) >= 2:
                    num = process_node(node.nodeargd.argnlist[0])
                    denom = process_node(node.nodeargd.argnlist[1])
                    return f' {num} divided by {denom} '
                return ' fraction '
            elif macro == 'sqrt':
                if node.nodeargd and hasattr(node.nodeargd, 'argnlist') and len(node.nodeargd.argnlist) > 0:
                    arg = process_node(node.nodeargd.argnlist[0])
                    return f' square root of {arg} '
                return ' square root '
            elif macro in ['cdot', 'times']:
                return ' times '
            elif macro == 'div':
                return ' divided by '
            elif macro == 'pm':
                return ' plus or minus '
            elif macro == 'mp':
                return ' minus or plus '
            elif macro == 'neq':
                return ' not equal to '
            elif macro == 'leq':
                return ' less than or equal to '
            elif macro == 'geq':
                return ' greater than or equal to '
            elif macro == 'approx':
                return ' approximately equal to '
            elif macro == 'equiv':
                return ' equivalent to '
            elif macro == 'propto':
                return ' proportional to '
            elif macro == 'infty':
                return ' infinity '
            elif macro == 'partial':
                return ' partial derivative '
            elif macro == 'nabla':
                return ' del '
            elif macro in ['to', 'rightarrow']:
                return ' approaches '
            elif macro == 'sin':
                return ' sine '
            elif macro == 'cos':
                return ' cosine '
            elif macro == 'tan':
                return ' tangent '
            elif macro == 'sec':
                return ' secant '
            elif macro == 'csc':
                return ' cosecant '
            elif macro == 'cot':
                return ' cotangent '
            elif macro == 'arcsin':
                return ' arcsine '
            elif macro == 'arccos':
                return ' arccosine '
            elif macro == 'arctan':
                return ' arctangent '
            elif macro == 'sinh':
                return ' hyperbolic sine '
            elif macro == 'cosh':
                return ' hyperbolic cosine '
            elif macro == 'tanh':
                return ' hyperbolic tangent '
            elif macro == 'log':
                return ' logarithm '
            elif macro == 'ln':
                return ' natural logarithm '
            elif macro == 'exp':
                return ' exponential '
            elif macro == 'max':
                return ' maximum '
            elif macro == 'min':
                return ' minimum '
            elif macro == 'sup':
                return ' supremum '
            elif macro == 'inf':
                return ' infimum '
            elif macro == 'det':
                return ' determinant '
            elif macro == 'gcd':
                return ' greatest common divisor '
            elif macro == 'arg':
                return ' argument '
            elif macro == 'deg':
                return ' degree '
            elif macro == 'dim':
                return ' dimension '
            elif macro == 'cup':
                return ' union '
            elif macro == 'cap':
                return ' intersection '
            elif macro == 'in':
                return ' element of '
            elif macro == 'notin':
                return ' not element of '
            elif macro in ['subset', 'subseteq']:
                return ' subset of '
            elif macro in ['supset', 'supseteq']:
                return ' superset of '
            elif macro == 'emptyset':
                return ' empty set '
            elif macro == 'forall':
                return ' for all '
            elif macro == 'exists':
                return ' there exists '
            elif macro == 'land':
                return ' logical and '
            elif macro == 'wedge':
                return ' wedge '
            elif macro == 'lor':
                return ' logical or '
            elif macro == 'vee':
                return ' vee '
            elif macro in ['neg', 'lnot']:
                return ' logical not '
            elif macro in ['implies', 'Rightarrow']:
                return ' implies '
            elif macro in ['iff', 'Leftrightarrow']:
                return ' biconditional '
            else:
                return f' {macro} '
        
        if hasattr(node, 'chars'):
            return node.chars
        
        return ''
    
    def process_latex(latex_expr):
        latex_expr = latex_expr.strip()
        
        chemical_patterns = [
            r'\b(NADH|NAD|ATP|ADP|DNA|RNA|CO2|H2O|NH3|CH4|FAD|FADH|CoA)\b',
            r'\b[A-Z][a-z]?\d*\^?[\+\-]?\d*\b',
        ]
        for pattern in chemical_patterns:
            if re.search(pattern, latex_expr):
                return latex_expr
        
        try:
            walker = latexwalker.LatexWalker(latex_expr, latex_context=lwc)
            nodes, _, _ = walker.get_latex_nodes()
            result = ''.join(process_node(node) for node in nodes)
            return result
        except:
            return latex_expr
    
    pattern = r'\$\$([^\$]+)\$\$|\$([^\$]+)\$'
    
    def handle_match(match):
        latex_expr = match.group(1) if match.group(1) else match.group(2)
        result = process_latex(latex_expr)
        return result
    
    processed_doc = re.sub(pattern, handle_match, doc)
    return processed_doc


def create_docs() -> None:
    """
    Creates document strings for question records to use with BERTopic.
    
    Pipeline:
    1. Extract text from images using OCR (if applicable)
    2. Generate image captions using BLIP
    3. Combine question text + OCR text + image captions into a single document string
    
    This creates rich textual representations of multimodal content that BERTopic
    can use for topic modeling while preserving academic vocabulary.
    
    Fetches records from DB, processes them, and saves back to DB until complete.
    """
    processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
    model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")

    db = initialize_and_fetch_db()
    processed_count = 0
    
    print("Starting doc creation process for BERTopic...")
    
    while True:
        # Get next record that needs doc creation
        record = get_empty_doc_record(db)
        
        # Break if no more records need processing
        if record is None:
            print(f"Doc creation complete! Processed {processed_count} records.")
            break
        
        try:
            # Simplify the record format
            record = simplify_question_record(record)
            
            question_id = record.get('question_id', '')
            question_text = record.get('question_text', '')
            answer_text = record.get('answer_text', '')
            question_media = record.get('question_media', [])
            answer_media = record.get('answer_media', [])
            
            # Helper function to safely convert tensor to PIL
            def safe_tensor_to_pil(img_tensor):
                if img_tensor is None:
                    return None
                
                # If it's already a PIL Image, return as-is
                if isinstance(img_tensor, Image.Image):
                    return img_tensor
                
                # Convert tensor to PIL safely
                if torch.is_tensor(img_tensor):
                    # Normalize tensor to [0,1] range
                    img_tensor = img_tensor.float()
                    img_tensor = (img_tensor - img_tensor.min()) / (img_tensor.max() - img_tensor.min())
                    
                    # Remove batch dimension if present
                    if img_tensor.dim() == 4:
                        img_tensor = img_tensor.squeeze(0)
                    
                    # Ensure correct channel order (C, H, W) -> (H, W, C)
                    if img_tensor.dim() == 3 and img_tensor.shape[0] in [1, 3]:
                        img_tensor = img_tensor.permute(1, 2, 0)
                    
                    # Convert to uint8 and PIL
                    img_array = (img_tensor * 255).clamp(0, 255).byte().numpy()
                    
                    # Handle grayscale
                    if img_array.shape[-1] == 1:
                        img_array = img_array.squeeze(-1)
                        return Image.fromarray(img_array, mode='L')
                    else:
                        return Image.fromarray(img_array, mode='RGB')
                
                return None
            
            # Process images to text descriptions
            image_descriptions = []
            ocr_texts = []
            
            # Process question media images
            if question_media:
                for img_name in question_media:
                    if img_name:
                        try:
                            img_tensor = load_image(img_name)
                            img_pil = safe_tensor_to_pil(img_tensor)
                            if img_pil is not None:
                                # Generate image caption using BLIP
                                inputs = processor(img_pil, return_tensors="pt")
                                with torch.no_grad():
                                    out = model.generate(**inputs, max_length=50, num_beams=5)
                                    caption = processor.decode(out[0], skip_special_tokens=True)
                                    image_descriptions.append(f"Question image: {caption}")
                                
                                # Extract text using OCR if image contains text
                                try:
                                    ocr_text = pytesseract.image_to_string(img_pil, config='--psm 6').strip()
                                    if ocr_text and len(ocr_text) > 3:  # Filter out noise
                                        ocr_texts.append(ocr_text) # only append the text, do not add random text to mark it.
                                except:
                                    pass  # OCR failed, continue without it
                        except Exception as e:
                            print(f"Error processing question image {img_name}: {e}")
            
            # Process answer media images
            if answer_media:
                for img_name in answer_media:
                    if img_name:
                        try:
                            img_tensor = load_image(img_name)
                            img_pil = safe_tensor_to_pil(img_tensor)
                            if img_pil is not None:
                                # Generate image caption using BLIP
                                inputs = processor(img_pil, return_tensors="pt")
                                with torch.no_grad():
                                    out = model.generate(**inputs, max_length=50, num_beams=5)
                                    caption = processor.decode(out[0], skip_special_tokens=True)
                                    image_descriptions.append(caption) # Additional explanatory text muddies the topic model
                                
                                # Extract text using OCR if image contains text
                                try:
                                    ocr_text = pytesseract.image_to_string(img_pil, config='--psm 6').strip()
                                    if ocr_text and len(ocr_text) > 3:  # Filter out noise
                                        ocr_texts.append(ocr_text)
                                except:
                                    pass  # OCR failed, continue without it
                        except Exception as e:
                            print(f"Error processing answer image {img_name}: {e}")
            
            # Combine all textual information into a single document
            all_text_components = []
            
            # Add original question and answer text
            if question_text:
                all_text_components.append(question_text)
            if answer_text:
                all_text_components.append(answer_text)
            
            # Add OCR extracted text
            if ocr_texts:
                all_text_components.extend(ocr_texts)
            
            # Add image descriptions
            if image_descriptions:
                all_text_components.extend(image_descriptions)
            
            # Create final doc string for BERTopic
            doc = " ".join(all_text_components)
            
            # Ensure we have some text
            if not doc.strip():
                doc = "a the is"
            
            # Update the record with the doc
            record['doc'] = convert_latex_to_plain_english(doc)
            
            # Save back to database
            success = upsert_question_record(db, record)
            
            if success:
                processed_count += 1
                print(f"Processed record {processed_count}: {question_id}")
                # print(f"  - Image descriptions: {len(image_descriptions)}")
                # print(f"  - OCR extractions: {len(ocr_texts)}")
                # print(f"  - Doc length: {len(doc)} chars")
            else:
                print(f"WARNING!!: Failed to save record: {question_id}")
            
        except Exception as e:
            print(f"Error processing record {record.get('question_id', 'unknown')}: {e}")
            
            # Still try to save a fallback record to avoid infinite loop
            try:
                record['doc'] = "Empty educational content"
                upsert_question_record(db, record)
                processed_count += 1
                print(f"Saved fallback for record: {record.get('question_id', 'unknown')}")
            except Exception as save_error:
                print(f"Failed to save fallback record: {save_error}")
                break  # Prevent infinite loop if DB is broken
    
    db.close()
    print("Database connection closed.")

def export_outlier_topics_to_docx(topic_model, output_path='outlier_topics.docx'):
    """
    Exports questions with topic_id = -1 to a .docx file with topic info.
    
    Args:
        topic_model: Fitted BERTopic model
        output_path: Output .docx file path
    """
    db = initialize_and_fetch_db()
    cursor = db.cursor()
    
    cursor.execute("""
        SELECT question_id, doc 
        FROM question_answer_pairs 
        WHERE topic_id = -1
    """)
    outliers = cursor.fetchall()
    
    if not outliers:
        print("No outlier questions found (topic_id = -1)")
        return
    
    doc = Document()
    
    topic_info = topic_model.get_topic(-1)
    
    header = doc.add_heading('Outlier Topic (-1)', level=1)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    keywords_para = doc.add_paragraph('Keywords: ')
    keywords_para.add_run(', '.join([word for word, _ in topic_info[:10]])).bold = True
    
    doc.add_paragraph('_' * 80)
    
    doc.add_heading('Representative Documents', level=2)
    rep_docs = topic_model.get_representative_docs(-1)
    for i, rep_doc in enumerate(rep_docs[:5], 1):
        doc.add_paragraph(f"{i}. {rep_doc[:200]}...")
    
    doc.add_paragraph('_' * 80)
    
    doc.add_heading('All Outlier Questions', level=2)
    for i, (question_id, doc_text) in enumerate(outliers, 1):
        doc.add_paragraph(f"Question {i} (ID: {question_id})")
        doc.add_paragraph(doc_text)
        doc.add_paragraph('_' * 80)
    
    doc.save(output_path)
    print(f"Exported {len(outliers)} outlier questions to {output_path}")

def set_process_limits():
    """
    Set CPU affinity and process priority for Linux systems.
    Limits the process to use only half the CPU cores and sets low priority.
    """
    if sys.platform != 'linux':
        return
    
    cpu_count = os.cpu_count()
    if cpu_count is None or cpu_count <= 1:
        return
    
    half_cores = max(1, cpu_count // 2)
    affinity_cores = list(range(half_cores))
    
    p = psutil.Process()
    p.cpu_affinity(affinity_cores)
    
    os.nice(10)
    
    print(f"Process limited to {half_cores} CPU cores (out of {cpu_count})")
    print(f"Process priority set to low (nice value increased by 10)")